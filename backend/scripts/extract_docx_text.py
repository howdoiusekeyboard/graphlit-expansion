"""Extract readable text from .docx files into a single Markdown file.

Usage (from repo root):
    python scripts/extract_docx_text.py --input-dir new --output new/docx_extracted.md
"""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document


def extract_docx_to_markdown(docx_path: Path) -> list[str]:
    """Extract paragraphs and tables from a DOCX file into Markdown lines."""
    doc = Document(str(docx_path))

    lines: list[str] = []
    lines.append("---")
    lines.append(f"## {docx_path.name}")
    lines.append("")

    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    if paragraphs:
        lines.append("### Paragraphs")
        lines.append("")
        for text in paragraphs:
            lines.append(text)
            lines.append("")

    if doc.tables:
        lines.append("### Tables")
        lines.append("")
        for t_i, table in enumerate(doc.tables, start=1):
            lines.append(f"**Table {t_i}**")
            for row in table.rows:
                cells = [c.text.strip().replace("\n", " ") for c in row.cells]
                lines.append(" | ".join(cells))
            lines.append("")

    return lines


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input-dir", type=str, default="new")
    parser.add_argument("--output", type=str, default="new/docx_extracted.md")
    args = parser.parse_args()

    input_dir = Path(args.input_dir)
    output_path = Path(args.output)

    if not input_dir.exists() or not input_dir.is_dir():
        raise SystemExit(f"Input directory not found: {input_dir}")

    docx_files = sorted(input_dir.glob("*.docx"))

    lines: list[str] = []
    lines.append("# Premidsem DOCX extraction")
    lines.append("")
    lines.append(f"Extracted from `{input_dir.as_posix()}`")
    lines.append("")

    if not docx_files:
        lines.append("_No .docx files found._")
        output_path.write_text("\n".join(lines), encoding="utf-8")
        return 0

    for p in docx_files:
        lines.extend(extract_docx_to_markdown(p))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text("\n".join(lines), encoding="utf-8")
    print(f"Wrote {len(docx_files)} .docx files into {output_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())



